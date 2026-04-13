"""
Parser Service — Extract structured data from PDF and DOCX resumes
"""

import re
from pathlib import Path
from typing import List, Optional

import pdfplumber
import structlog
from docx import Document

from models.resume_model import (
    ParsedResumeData, ContactInfo, WorkExperience,
    Education, Project, Certification
)
from utils.nlp_utils import (
    clean_text, detect_skills_in_text, extract_email,
    extract_phone, extract_urls, count_words,
    extract_years_of_experience, extract_sections
)

logger = structlog.get_logger(__name__)


class ParserService:
    """Parses resumes from PDF/DOCX into structured ParsedResumeData."""

    # ── Public Entry Point ─────────────────────────────────────────────────────
    async def parse_resume(self, file_path: str, file_type: str) -> ParsedResumeData:
        logger.info("Parsing resume", file_path=file_path, file_type=file_type)
        try:
            raw_text = await self._extract_raw_text(file_path, file_type)
            return await self._structure_resume(raw_text)
        except Exception as e:
            logger.error("Resume parse failed", error=str(e))
            raise RuntimeError(f"Failed to parse resume: {str(e)}")

    # ── Text Extraction ────────────────────────────────────────────────────────
    async def _extract_raw_text(self, file_path: str, file_type: str) -> str:
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(self, file_path: str) -> str:
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            raise RuntimeError(f"PDF extraction error: {e}")
        full_text = "\n".join(text_parts)
        if not full_text.strip():
            raise ValueError("PDF contains no extractable text (may be image-based).")
        return full_text

    def _extract_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"DOCX extraction error: {e}")

    # ── Structure Parsing ─────────────────────────────────────────────────────
    async def _structure_resume(self, raw_text: str) -> ParsedResumeData:
        sections = extract_sections(raw_text)
        tech_skills, soft_skills = detect_skills_in_text(raw_text)
        contact = self._extract_contact(raw_text)
        full_name = self._extract_name(raw_text)
        experience = self._extract_experience(sections.get("experience", ""))
        education = self._extract_education(sections.get("education", ""))
        projects = self._extract_projects(sections.get("projects", ""))
        certs = self._extract_certifications(sections.get("certifications", ""))
        total_exp = extract_years_of_experience(raw_text)

        # Fallback: calculate experience from work entries
        if total_exp == 0.0 and experience:
            total_exp = self._calculate_experience_from_entries(experience)

        all_skills = list(set(tech_skills + soft_skills))

        return ParsedResumeData(
            raw_text=raw_text,
            contact_info=contact,
            full_name=full_name,
            summary=sections.get("summary", ""),
            skills=all_skills,
            technical_skills=tech_skills,
            soft_skills=soft_skills,
            work_experience=experience,
            education=education,
            projects=projects,
            certifications=certs,
            total_experience_years=total_exp,
            word_count=count_words(raw_text),
            sections_detected=list(sections.keys()),
        )

    def _extract_contact(self, text: str) -> ContactInfo:
        email = extract_email(text)
        phone = extract_phone(text)
        urls = extract_urls(text)
        linkedin = next((u for u in urls if "linkedin" in u.lower()), None)
        github = next((u for u in urls if "github" in u.lower()), None)
        portfolio = next(
            (u for u in urls if "linkedin" not in u.lower() and "github" not in u.lower()), None
        )
        location = self._extract_location(text)
        return ContactInfo(
            email=email, phone=phone, location=location,
            linkedin=linkedin, github=github, portfolio=portfolio,
        )

    def _extract_name(self, text: str) -> Optional[str]:
        """Heuristic: first non-empty line that looks like a name."""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines[:5]:
            words = line.split()
            if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w.isalpha()):
                if not re.search(r"[@\d\|]", line):
                    return line
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        patterns = [
            r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)?,\s*[A-Z]{2}(?:\s\d{5})?)\b",
            r"\b([A-Z][a-z]+,\s*[A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\b",
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return None

    def _extract_experience(self, text: str) -> List[WorkExperience]:
        if not text:
            return []
        experiences = []
        # Split on company/role boundaries (simple heuristic)
        blocks = re.split(r"\n(?=[A-Z][^\n]{5,60}\n)", text)
        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if len(lines) < 2:
                continue
            title_line = lines[0]
            company_line = lines[1] if len(lines) > 1 else ""
            desc = " ".join(lines[2:]) if len(lines) > 2 else ""

            # Extract date range
            date_pattern = r"(\w+\s*\d{4})\s*[-–—]\s*(\w+\s*\d{4}|present|current)"
            date_match = re.search(date_pattern, block, re.IGNORECASE)
            start_date = date_match.group(1) if date_match else None
            end_date = date_match.group(2) if date_match else None
            is_current = end_date and end_date.lower() in ("present", "current") if end_date else False

            # Skills in this block
            tech, _ = detect_skills_in_text(block)

            exp = WorkExperience(
                company=company_line or "Unknown",
                title=title_line,
                start_date=start_date,
                end_date=end_date,
                description=desc[:500],
                technologies=tech[:10],
                is_current=is_current,
            )
            experiences.append(exp)
        return experiences[:10]  # Cap at 10 entries

    def _extract_education(self, text: str) -> List[Education]:
        if not text:
            return []
        educations = []
        degree_patterns = [
            r"\b(B\.?S\.?|B\.?A\.?|M\.?S\.?|M\.?A\.?|Ph\.?D\.?|MBA|Bachelor|Master|Doctor)\b",
        ]
        year_pattern = r"\b(19|20)\d{2}\b"
        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            institution = lines[0]
            degree_match = None
            for pat in degree_patterns:
                degree_match = re.search(pat, block, re.IGNORECASE)
                if degree_match:
                    break
            years = re.findall(year_pattern, block)
            start_year = int("".join(years[0])) if years else None
            end_year = int("".join(years[1])) if len(years) > 1 else None

            gpa_match = re.search(r"GPA:?\s*(\d+\.\d+)", block, re.IGNORECASE)
            gpa = float(gpa_match.group(1)) if gpa_match else None

            educations.append(Education(
                institution=institution,
                degree=degree_match.group(0) if degree_match else None,
                start_year=start_year,
                end_year=end_year,
                gpa=gpa,
            ))
        return educations[:5]

    def _extract_projects(self, text: str) -> List[Project]:
        if not text:
            return []
        projects = []
        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            lines = [l.strip() for l in block.split("\n") if l.strip()]
            if not lines:
                continue
            name = lines[0]
            desc = " ".join(lines[1:])[:400]
            tech, _ = detect_skills_in_text(block)
            urls = extract_urls(block)
            github_url = next((u for u in urls if "github" in u), None)
            url = next((u for u in urls if u != github_url), None)
            projects.append(Project(
                name=name, description=desc, technologies=tech[:8],
                url=url, github_url=github_url,
            ))
        return projects[:10]

    def _extract_certifications(self, text: str) -> List[Certification]:
        if not text:
            return []
        certs = []
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines:
            if len(line) > 10:
                year_match = re.search(r"\b(20\d{2})\b", line)
                certs.append(Certification(
                    name=line[:200],
                    issue_date=year_match.group(1) if year_match else None,
                ))
        return certs[:10]

    def _calculate_experience_from_entries(self, experience: List[WorkExperience]) -> float:
        total_months = 0
        for exp in experience:
            if exp.duration_months:
                total_months += exp.duration_months
        return round(total_months / 12, 1)
